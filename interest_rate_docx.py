from datetime import datetime
from pathlib import Path
from docx import Document
from docx.document import Document as DocumentObject    # 객체

from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
import pandas as pd

from utils.api_data import OUT_DIR, IMG_DIR, OUT1, OUT2
from utils.init_docx import OUT3, apply_font

OUT4 = OUT_DIR / f"{Path(__file__).stem}.docx"

def add_blank_paragraph(doc: DocumentObject, size_pt: int=None):
    r_empty = doc.add_paragraph().add_run(" ")    # 빈 단락
    apply_font(r_empty, size_pt=size_pt)    # 폰트

def add_title(OUT3):
    doc = Document(OUT3)
    p_title = doc.add_paragraph(style="Title")    # 제목
    r_title = p_title.add_run("정기예금 금리 현황표")
    apply_font(r_title, face="Malgun Gothic", size_pt=20, is_bold=True)    # 폰트
    now = datetime.now()    # 현재 시점
    now_string = now.isoformat(sep=" ", timespec="minutes")    # 작성일시
    r_now = p_title.add_run(f" (작성 일시: {now_string})")
    apply_font(r_now, size_pt=14)
    add_blank_paragraph(doc, size_pt=6)    # 빈 단락
    doc.save(OUT4)



#######################################################################################

def insert_indicators(OUT4):
    doc = Document(OUT4)
    r_head = doc.add_paragraph().add_run("1. 주요 금리(최근 24개월)")
    apply_font(r_head, size_pt=14, is_bold=True)    # 폰트
    add_blank_paragraph(doc, size_pt=10)    # 빈 단락
    table = doc.add_table(rows=1, cols=5)     #  표 추가(1행, 5열)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER    # 가로 정렬
    # table.allow_autofit = False    # 자동 너비 맞춤 해제
    table.autofit = False

    tr = table.rows[0]
    with pd.ExcelFile(OUT2) as xlsx:
        for idx, sheet_name in enumerate(xlsx.sheet_names):
            df = pd.read_excel(xlsx, sheet_name=sheet_name)
            df = df.tail(36)    # 최근 3년 데이터
            td = tr.cells[idx]
            td.width = Mm(35.5)    # 셀 너비

            p1 = td.paragraphs[0]
            r1 = p1.add_run(sheet_name)    # 금리 지표
            apply_font(r1, size_pt=12, is_bold=True, rgb="333333")

            p2 = td.add_paragraph()
            last_value = df["DATA_VALUE"].iloc[-1]
            r2 = p2.add_run(f"{last_value:,.2f}")    # 전월말 데이터
            apply_font(r2, size_pt=14, is_bold=True, rgb="333333")

            p3 = td.add_paragraph()    # 단락 추가
            diff = last_value - df["DATA_VALUE"].iloc[0]
            arrow = "▲" if diff > 0 else "▼" if diff < 0 else ""
            color = ("FF0000" if diff > 0 else "0000FF" if diff < 0 else "000000")    # 색상
            r3 = p3.add_run(f"{arrow}{abs(diff):,.2f}%p")    # 지표 변화량
            apply_font(r3, size_pt=10, is_bold=True, rgb=color)

            p4 = td.add_paragraph()
            p4.paragraph_format.left_indent = Mm(-1)    # 들여쓰기
            img_path = IMG_DIR / f"{sheet_name}.png"
            p4.add_run().add_picture(img_path.as_posix(), Mm(30), Mm(8))    # 이미지
            
        add_blank_paragraph(doc, size_pt=10)
        doc.save(OUT4)   

###################################################################################################
OUT6 = OUT_DIR / "result.docx"

# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn

# def set_cell_width(cell, width):
#     """
#     테이블 셀의 너비를 고정 (dxa 단위)
#     :param cell: docx.table._Cell 객체
#     :param width: docx.shared.Mm 객체
#     """
#     cell.width = width
#     tc_pr = cell._tc.get_or_add_tcPr()

#     tc_w = OxmlElement("w:tcW")
#     tc_w.set(qn("w:type"), "dxa")
#     tc_w.set(qn("w:w"), str(int(width)))  # Mm → EMU 단위

#     # 중복 방지: 기존 tcW 제거
#     for child in tc_pr.findall(qn("w:tcW")):
#         tc_pr.remove(child)

#     tc_pr.append(tc_w)

# def set_table_width(table, widths):
#     """
#     테이블 전체 너비를 고정 (열 너비 리스트 기준)
#     :param table: docx.table.Table 객체
#     :param widths: 각 열의 docx.shared.Mm 리스트
#     """
#     tbl = table._tbl
#     tbl_pr = tbl.tblPr

#     tbl_w = OxmlElement("w:tblW")
#     tbl_w.set(qn("w:type"), "dxa")
#     tbl_w.set(qn("w:w"), str(sum(int(w) for w in widths)))  # 전체 너비 = 열 너비 합계

#     # 중복 방지: 기존 tblW 제거
#     for child in tbl_pr.findall(qn("w:tblW")):
#         tbl_pr.remove(child)

#     tbl_pr.append(tbl_w)


def insert_deposit_info(n_rows: int = 10):
    doc = Document(OUT4)
    r_head = doc.add_paragraph().add_run("2. 주요 정기예금 상품 및 금리")
    apply_font(r_head, size_pt=14, is_bold=True)
    add_blank_paragraph(doc, size_pt=2)

    table = doc.add_table(rows=1, cols=6, style="Light Shading Accent 4") 
    table.alignment = WD_TABLE_ALIGNMENT.CENTER    # 표 가로 정렬
    # table.allow_autofit = False    # 표 너비 자동 맞춤 해제
    table.autofit = False

    tr = table.rows[0]
    th_text = ['금융기관', '상품명', '이자계산', '만기(월)', '세전금리', '최고우대']
    col_width = [Mm(45), Mm(55), Mm(20), Mm(20), Mm(20), Mm(20)]    # 열 너비
    # # 테이블 전체 고정
    # set_table_width(table, col_width)
    # # 테이블 너비 자동조정 완전 해제
    # tbl_pr = table._tbl.tblPr
    # tbl_layout = OxmlElement("w:tblLayout")
    # tbl_layout.set(qn("w:type"), "fixed")
    # tbl_pr.append(tbl_layout)

    for idx, th in enumerate(tr.cells):
        # set_cell_width(th, col_width[idx])
        th.width = col_width[idx]
        th.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_th = th.paragraphs[0]
        p_th.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_th = p_th.add_run(f"{th_text[idx]}")
        apply_font(r_th, size_pt=12, is_bold=True)

    df = pd.read_excel(OUT1)
    df_filter = df.filter(["kor_co_nm", "fin_prdt_nm", "intr_rate_type_nm", "save_trm", "intr_rate", "intr_rate2"])
    df_sort = df_filter.sort_values("intr_rate", ascending=False)    # 금리를 기준으로 상위 10개만

    for _, se_row in df_sort.head(n_rows).iterrows():     # n_rows에 지정된 개수 반복 출력(상위 10개)
        tr = table.add_row()    # 행 추가
        for idx, td in enumerate(tr.cells):
            td.width = col_width[idx]
            td.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_td = td.paragraphs[0]
            if idx < 2:    # 첫 두개의 셀은 텍스트 배분 정렬
                p_td.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
            p_td.paragraph_format.space_before = Mm(2)
            p_td.paragraph_format.space_after = Mm(2)
            p_td.add_run(f"{se_row.iloc[idx]}")    # 내용 셀 입력
    
    add_blank_paragraph(doc, size_pt=10)
    doc.save(OUT6)

OUT_X = OUT_DIR / "mini_result.docx"

def insert_info(OUT6):
    doc = Document(OUT6)
    table = doc.add_table(rows=1, cols=1, style="Light Shading Accent 6")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tr = table.rows[0]
    td = tr.cells[0]
    td.width = Mm(174)

    p1 = td.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run("주의사항")
    apply_font(r1, size_pt=12, is_bold=True)

    p2 = td.add_paragraph(style="List Bullet")
    r2 = p2.add_run("이번 장에서 작성한 정기예금 금리 현황표는 API 호출 시점에 따라 값이 다르므로, 참고용으로만 사용하세요.")
    apply_font(r2, size_pt=10, rgb="f79646")

    p3 = td.add_paragraph(style="List Bullet")
    r3 = p3.add_run("정기예금 상품의 금리는 수시로 변경될 수 있으므로, 거래전 반드시 해당 금융회사에 문의하시기 바랍니다.")
    apply_font(r3, size_pt=10, rgb="f79646")

    add_blank_paragraph(doc, size_pt=10)
    doc.save(OUT_X)


if __name__ == "__main__":
    add_title(OUT3)
    insert_indicators(OUT4)
    insert_deposit_info(10)
    insert_info(OUT6)